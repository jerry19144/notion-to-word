"""
Word document formatting and styling utilities
"""

from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)


class WordFormatter:
    """Handles Word document formatting and styling"""

    def __init__(self, template_config: Dict[str, Any]):
        """Initialize formatter with template configuration"""
        self.config = template_config
        self.font = self.config.get("font", "Calibri")
        self.formatting = self.config.get("formatting", {})

    def create_document(self, template_path: Optional[str] = None) -> Document:
        """Create a new Word document from template or blank"""
        if template_path:
            try:
                doc = Document(template_path)
                # Clear template content while preserving styles
                for paragraph in list(doc.paragraphs):
                    p = paragraph._element
                    p.getparent().remove(p)
                for table in list(doc.tables):
                    t = table._element
                    t.getparent().remove(t)
                logger.info(f"Created document from template: {template_path}")
                return doc
            except Exception as e:
                logger.warning(f"Could not load template: {e}, using blank document")

        return Document()

    def apply_style_safe(self, paragraph, style_name: str) -> bool:
        """Safely apply a style to a paragraph with fallback"""
        try:
            paragraph.style = style_name
            return True
        except:
            try:
                paragraph.style = "Normal"
                logger.debug(f"Style '{style_name}' not found, using 'Normal'")
                return True
            except:
                logger.debug(f"Could not apply any style to paragraph")
                return False

    def create_paragraph_safe(self, doc: Document, text: str = "", style: str = "Normal"):
        """Safely create a paragraph with fallback style handling"""
        try:
            p = doc.add_paragraph(text, style=style)
            return p
        except:
            try:
                p = doc.add_paragraph(text, style="Normal")
                return p
            except:
                p = doc.add_paragraph(text)
                return p

    def apply_rich_text_formatting(self, paragraph, rich_text_array: List[Dict]):
        """Apply formatting to rich text in a paragraph"""
        for text_block in rich_text_array:
            if text_block['type'] != 'text':
                continue

            content = text_block['text']['content']
            annotations = text_block.get('annotations', {})

            run = paragraph.add_run(content)
            run.font.name = self.font

            # Apply formatting based on annotations
            if annotations.get('bold'):
                run.bold = True
            if annotations.get('italic'):
                run.italic = True
            if annotations.get('strikethrough'):
                run.font.strike = True
            if annotations.get('underline'):
                run.underline = True
            if annotations.get('code'):
                code_font = self.formatting.get("code_font", "Courier New")
                run.font.name = code_font
                run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

            # Handle color
            if annotations.get('color') and annotations['color'] != 'default':
                color_map = {
                    'gray': RGBColor(0x80, 0x80, 0x80),
                    'brown': RGBColor(0xA5, 0x2A, 0x2A),
                    'orange': RGBColor(0xFF, 0xA5, 0x00),
                    'yellow': RGBColor(0xFF, 0xFF, 0x00),
                    'green': RGBColor(0x00, 0x80, 0x00),
                    'blue': RGBColor(0x00, 0x00, 0xFF),
                    'purple': RGBColor(0x80, 0x00, 0x80),
                    'pink': RGBColor(0xFF, 0xC0, 0xCB),
                    'red': RGBColor(0xFF, 0x00, 0x00),
                }
                if annotations['color'] in color_map:
                    run.font.color.rgb = color_map[annotations['color']]

    def add_title(self, doc: Document, title: str, style: str = "Title"):
        """Add a formatted title to the document"""
        title_para = self.create_paragraph_safe(doc, title, style)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply title formatting
        title_size = self.formatting.get("title_size", 16)
        for run in title_para.runs:
            run.font.name = self.font
            run.font.size = Pt(title_size)
            run.bold = True

        return title_para

    def add_metadata(self, doc: Document, created_time: str, last_edited: str):
        """Add document metadata (creation and edit times)"""
        metadata_text = ""
        if created_time:
            metadata_text += f"Created: {created_time[:10]}"
        if last_edited:
            if metadata_text:
                metadata_text += "  |  "
            metadata_text += f"Last edited: {last_edited[:10]}"

        if metadata_text:
            p = self.create_paragraph_safe(doc, metadata_text, "Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            metadata_size = self.formatting.get("metadata_size", 10)
            metadata_color = self.formatting.get("metadata_color", "#666666")

            # Convert hex color to RGB
            if metadata_color.startswith("#"):
                hex_color = metadata_color.lstrip("#")
                rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                color = RGBColor(*rgb)
            else:
                color = RGBColor(0x66, 0x66, 0x66)

            for run in p.runs:
                run.font.size = Pt(metadata_size)
                run.font.color.rgb = color
                run.font.name = self.font

            # Add spacing after metadata
            doc.add_paragraph()

    def add_divider(self, doc: Document, style: str = "Normal"):
        """Add a divider line to the document"""
        p = self.create_paragraph_safe(doc, "_" * 50, style)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)