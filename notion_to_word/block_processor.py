"""
Process different types of Notion blocks
"""

from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import logging
import requests

from .notion_client import NotionClient
from .word_formatter import WordFormatter

logger = logging.getLogger(__name__)


class BlockProcessor:
    """Process Notion blocks and convert them to Word format"""

    def __init__(self, notion_client: NotionClient, formatter: WordFormatter, config: Dict):
        """Initialize block processor"""
        self.notion = notion_client
        self.formatter = formatter
        self.config = config
        self.formatting = config.get("formatting", {})
        self.stats = {
            "blocks_processed": 0,
            "blocks_with_errors": 0,
            "block_types_seen": set()
        }

    def process_blocks(self, blocks: List[Dict], doc: Document, style_mappings: Dict, level: int = 0):
        """Process all blocks recursively"""
        for block in blocks:
            try:
                block_type = block['type']
                self.stats["block_types_seen"].add(block_type)
                self.stats["blocks_processed"] += 1

                # Get appropriate style from mappings
                style_name = style_mappings.get(block_type, style_mappings.get("default", "Normal"))

                # Process based on type
                if block_type == 'paragraph':
                    self._add_paragraph(doc, block, level, style_name)
                elif block_type.startswith('heading_'):
                    self._add_heading(doc, block, block_type, style_name)
                elif block_type == 'bulleted_list_item':
                    self._add_bullet_item(doc, block, level, style_name)
                elif block_type == 'numbered_list_item':
                    self._add_numbered_item(doc, block, level, style_name)
                elif block_type == 'to_do':
                    self._add_todo_item(doc, block, level, style_name)
                elif block_type == 'toggle':
                    self._add_toggle(doc, block, level, style_name)
                elif block_type == 'code':
                    self._add_code_block(doc, block, style_name)
                elif block_type == 'quote':
                    self._add_quote(doc, block, style_name)
                elif block_type == 'callout':
                    self._add_callout(doc, block, style_name)
                elif block_type == 'divider':
                    self.formatter.add_divider(doc, style_name)
                elif block_type == 'image':
                    self._add_image(doc, block, style_name)
                elif block_type == 'table':
                    self._add_table(doc, block)
                elif block_type == 'table_row':
                    pass  # handled inside _add_table
                else:
                    self._add_unknown_block(doc, block, style_name)

                # Process children if they exist (table blocks handle their own rows)
                if block.get('has_children') and block_type not in ('table', 'table_row'):
                    children = self.notion.fetch_children_blocks(block['id'])
                    if children:
                        self.process_blocks(children, doc, style_mappings, level + 1)

            except Exception as e:
                self.stats["blocks_with_errors"] += 1
                logger.error(f"Error processing block {block.get('id', 'unknown')}: {e}")
                self._add_error_placeholder(doc)

    def _add_paragraph(self, doc: Document, block: Dict, level: int, style: str):
        """Add a paragraph block"""
        paragraph_data = block.get('paragraph', {})
        rich_text = paragraph_data.get('rich_text', [])

        if rich_text:
            p = self.formatter.create_paragraph_safe(doc, style=style)
            if level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * level)
            self.formatter.apply_rich_text_formatting(p, rich_text)
        else:
            doc.add_paragraph()

    def _add_heading(self, doc: Document, block: Dict, heading_type: str, style: str):
        """Add a heading block"""
        heading_data = block.get(heading_type, {})
        rich_text = heading_data.get('rich_text', [])
        text = NotionClient.extract_text_from_rich_text(rich_text)

        p = self.formatter.create_paragraph_safe(doc, text, style)

        # If falling back to Normal, make it bold and larger
        if "Normal" in p.style.name:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(14)

    def _add_bullet_item(self, doc: Document, block: Dict, level: int, style: str):
        """Add a bulleted list item"""
        list_data = block.get('bulleted_list_item', {})
        rich_text = list_data.get('rich_text', [])

        bullet = self.formatting.get("bullet_symbol", "•")
        p = self.formatter.create_paragraph_safe(doc, style=style)

        if level > 0:
            p.paragraph_format.left_indent = Inches(0.25 * level)

        p.add_run(f"{bullet} ")
        self.formatter.apply_rich_text_formatting(p, rich_text)

    def _add_numbered_item(self, doc: Document, block: Dict, level: int, style: str):
        """Add a numbered list item"""
        list_data = block.get('numbered_list_item', {})
        rich_text = list_data.get('rich_text', [])

        p = self.formatter.create_paragraph_safe(doc, style=style)
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.25 * level)
        self.formatter.apply_rich_text_formatting(p, rich_text)

    def _add_todo_item(self, doc: Document, block: Dict, level: int, style: str):
        """Add a to-do item"""
        todo_data = block.get('to_do', {})
        rich_text = todo_data.get('rich_text', [])
        checked = todo_data.get('checked', False)

        checkbox = self.formatting.get("checkbox_checked", "☑") if checked else self.formatting.get("checkbox_unchecked", "☐")

        p = self.formatter.create_paragraph_safe(doc, style=style)
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.25 * level)
        p.add_run(f"{checkbox} ")
        self.formatter.apply_rich_text_formatting(p, rich_text)

    def _add_toggle(self, doc: Document, block: Dict, level: int, style: str):
        """Add a toggle block"""
        toggle_data = block.get('toggle', {})
        rich_text = toggle_data.get('rich_text', [])

        toggle_symbol = self.formatting.get("toggle_symbol", "▶")
        p = self.formatter.create_paragraph_safe(doc, style=style)
        if level > 0:
            p.paragraph_format.left_indent = Inches(0.25 * level)
        p.add_run(f"{toggle_symbol} ")
        self.formatter.apply_rich_text_formatting(p, rich_text)

    def _add_code_block(self, doc: Document, block: Dict, style: str):
        """Add a code block"""
        code_data = block.get('code', {})
        rich_text = code_data.get('rich_text', [])
        language = code_data.get('language', 'plain text')
        text = NotionClient.extract_text_from_rich_text(rich_text)

        # Language label
        p = self.formatter.create_paragraph_safe(doc, f"[{language}]", style)
        p.runs[0].italic = True

        # Code content
        p = self.formatter.create_paragraph_safe(doc, text, style)
        p.paragraph_format.left_indent = Inches(0.25)
        code_font = self.formatting.get("code_font", "Courier New")
        code_size = self.formatting.get("code_size", 10)
        for run in p.runs:
            run.font.name = code_font
            run.font.size = Pt(code_size)

    def _add_quote(self, doc: Document, block: Dict, style: str):
        """Add a quote block"""
        quote_data = block.get('quote', {})
        rich_text = quote_data.get('rich_text', [])

        quote_left = self.formatting.get("quote_left", "\u201c")
        quote_right = self.formatting.get("quote_right", "\u201d")

        p = self.formatter.create_paragraph_safe(doc, style=style)
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(quote_left)
        self.formatter.apply_rich_text_formatting(p, rich_text)
        p.add_run(quote_right)

        for run in p.runs:
            run.italic = True

    def _add_callout(self, doc: Document, block: Dict, style: str):
        """Add a callout block"""
        callout_data = block.get('callout', {})
        rich_text = callout_data.get('rich_text', [])
        icon = callout_data.get('icon', {})

        p = self.formatter.create_paragraph_safe(doc, style=style)
        p.paragraph_format.left_indent = Inches(0.25)

        if icon and icon.get('type') == 'emoji':
            p.add_run(f"{icon.get('emoji', '💡')} ")

        self.formatter.apply_rich_text_formatting(p, rich_text)

    def _add_image(self, doc: Document, block: Dict, style: str):
        """Download and embed an image from a Notion image block"""
        image_data = block.get('image', {})
        image_type = image_data.get('type')

        url = None
        if image_type == 'file':
            url = image_data.get('file', {}).get('url')
        elif image_type == 'external':
            url = image_data.get('external', {}).get('url')

        embedded = False
        if url:
            try:
                response = requests.get(url, timeout=30)
                response.raise_for_status()
                image_stream = io.BytesIO(response.content)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(image_stream, width=Inches(6))
                embedded = True
            except Exception as e:
                logger.warning(f"Could not embed image from {url}: {e}")

        if not embedded:
            p = self.formatter.create_paragraph_safe(doc, "[Image could not be loaded]", style)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].italic = True

        # Add caption if present
        caption = image_data.get('caption', [])
        if caption:
            text = NotionClient.extract_text_from_rich_text(caption)
            cap_p = self.formatter.create_paragraph_safe(doc, text, style)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap_p.runs:
                run.font.size = Pt(10)
                run.italic = True

    def _add_table(self, doc: Document, block: Dict):
        """Build a Word table from a Notion table block"""
        table_data = block.get('table', {})
        table_width = table_data.get('table_width', 1)
        has_column_header = table_data.get('has_column_header', False)

        rows = self.notion.fetch_children_blocks(block['id'])
        if not rows:
            return

        word_table = doc.add_table(rows=len(rows), cols=table_width)
        word_table.style = 'Table Grid'

        for row_idx, row_block in enumerate(rows):
            if row_block.get('type') != 'table_row':
                continue
            cells = row_block.get('table_row', {}).get('cells', [])
            for col_idx, cell_rich_text in enumerate(cells):
                if col_idx >= table_width:
                    break
                cell_para = word_table.cell(row_idx, col_idx).paragraphs[0]
                if cell_rich_text:
                    self.formatter.apply_rich_text_formatting(cell_para, cell_rich_text)
                if has_column_header and row_idx == 0:
                    for run in cell_para.runs:
                        run.bold = True

        doc.add_paragraph()

    def _add_unknown_block(self, doc: Document, block: Dict, style: str):
        """Handle unknown block types"""
        block_type = block.get('type', 'unknown')
        p = self.formatter.create_paragraph_safe(doc, f"[{block_type} block]", style)
        if p.runs:
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def _add_error_placeholder(self, doc: Document):
        """Add an error placeholder"""
        try:
            p = self.formatter.create_paragraph_safe(doc, "[Error processing block]", "Normal")
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        except:
            pass