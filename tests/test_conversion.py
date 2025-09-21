#!/usr/bin/env python3
"""
Test actual conversion with a real Notion page
"""

import os
import sys
import time
from docx import Document

from notion_to_word.converter import NotionToWordConverter


def test_real_conversion():
    """Test with a real Notion page URL"""

    # You can change this to any test page URL
    TEST_URL = "https://www.notion.so/nexuslux/Icebreaker-Global-Visit-2025-DG-Draft-274e531aede680e9a2a7c31931ae355e?source=copy_link"

    print("\n" + "="*60)
    print(" REAL CONVERSION TEST")
    print("="*60)

    # Load token
    from dotenv import load_dotenv
    load_dotenv()
    token = os.getenv('NOTION_TOKEN')

    if not token:
        print("\n❌ No NOTION_TOKEN found in .env")
        return False

    # Initialize converter
    print(f"\n📝 Testing with: {TEST_URL[:60]}...")

    converter = NotionToWordConverter(
        token,
        template_path="IC Template.docx" if os.path.exists("IC Template.docx") else None,
        output_dir="outputs"
    )

    # Measure conversion time
    start_time = time.time()

    try:
        # Convert the page
        output_path = converter.convert_page(TEST_URL)

        end_time = time.time()
        duration = end_time - start_time

        # Verify output
        if not os.path.exists(output_path):
            print("❌ Output file not created")
            return False

        # Check file size
        file_size = os.path.getsize(output_path)

        # Load and verify document content
        doc = Document(output_path)
        para_count = len(doc.paragraphs)

        # Get some content preview
        content_preview = ""
        for i, para in enumerate(doc.paragraphs[:3]):
            if para.text.strip():
                content_preview += f"  {i+1}. {para.text[:60]}...\n"

        # Print results
        print(f"\n✅ Conversion successful!")
        print(f"   Time: {duration:.2f} seconds")
        print(f"   File: {output_path}")
        print(f"   Size: {file_size:,} bytes")
        print(f"   Paragraphs: {para_count}")

        if content_preview:
            print(f"\n📄 Content preview:")
            print(content_preview)

        # Performance assessment
        if duration < 5:
            print(f"\n⚡ Performance: Excellent ({duration:.2f}s)")
        elif duration < 10:
            print(f"\n✓ Performance: Good ({duration:.2f}s)")
        else:
            print(f"\n⚠️  Performance: Slow ({duration:.2f}s)")

        # Check for processing stats
        if converter.processor.stats["blocks_with_errors"] > 0:
            print(f"\n⚠️  Blocks with errors: {converter.processor.stats['blocks_with_errors']}")

        return True

    except Exception as e:
        print(f"\n❌ Conversion failed: {e}")
        return False


def test_batch_conversion():
    """Test converting multiple pages quickly"""
    print("\n" + "="*60)
    print(" BATCH CONVERSION TEST")
    print("="*60)

    # Test URLs - add more if you have them
    test_urls = [
        "274e531aede680e9a2a7c31931ae355e",  # Just ID
        # Add more test URLs here
    ]

    from dotenv import load_dotenv
    load_dotenv()
    token = os.getenv('NOTION_TOKEN')

    if not token:
        print("\n❌ No token for batch test")
        return

    converter = NotionToWordConverter(token, output_dir="outputs")

    print(f"\n📚 Converting {len(test_urls)} pages...")

    results = []
    total_time = 0

    for i, url in enumerate(test_urls, 1):
        try:
            print(f"\n[{i}/{len(test_urls)}] Converting...")
            start = time.time()
            output = converter.convert_page(url)
            duration = time.time() - start
            total_time += duration

            results.append({
                "url": url[:40],
                "output": output,
                "time": duration,
                "success": True
            })
            print(f"   ✓ Done in {duration:.2f}s")

        except Exception as e:
            results.append({
                "url": url[:40],
                "error": str(e),
                "success": False
            })
            print(f"   ✗ Failed: {e}")

    # Summary
    successful = sum(1 for r in results if r["success"])
    print(f"\n📊 Batch Results:")
    print(f"   Total: {len(results)} pages")
    print(f"   Successful: {successful}")
    print(f"   Failed: {len(results) - successful}")
    if successful > 0:
        avg_time = total_time / successful
        print(f"   Average time: {avg_time:.2f}s per page")


def main():
    """Run all conversion tests"""

    # Test single conversion
    success = test_real_conversion()

    # Optional: Test batch conversion
    # test_batch_conversion()

    if success:
        print("\n✅ All conversion tests passed!")
        return 0
    else:
        print("\n❌ Some tests failed")
        return 1


if __name__ == "__main__":
    exit_code = main()
    sys.exit(exit_code)