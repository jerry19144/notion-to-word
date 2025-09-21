#!/usr/bin/env python3
"""
Comprehensive test suite for Notion to Word converter
Run this to test all functionality
"""

import os
import sys
import json
import time
from datetime import datetime
from typing import List, Dict, Any
from docx import Document

from notion_to_word.notion_client import NotionClient
from notion_to_word.config import TemplateConfig
from notion_to_word.word_formatter import WordFormatter
from notion_to_word.converter import NotionToWordConverter

# Test results storage
test_results = {
    "passed": [],
    "failed": [],
    "skipped": []
}


class Colors:
    """Console colors for output"""
    GREEN = '\033[92m'
    RED = '\033[91m'
    YELLOW = '\033[93m'
    BLUE = '\033[94m'
    RESET = '\033[0m'
    BOLD = '\033[1m'


def print_header(text: str):
    """Print a test section header"""
    print(f"\n{Colors.BLUE}{Colors.BOLD}{'=' * 60}")
    print(f" {text}")
    print(f"{'=' * 60}{Colors.RESET}")


def print_test(name: str, status: str, message: str = ""):
    """Print test result"""
    if status == "PASS":
        color = Colors.GREEN
        symbol = "✓"
        test_results["passed"].append(name)
    elif status == "FAIL":
        color = Colors.RED
        symbol = "✗"
        test_results["failed"].append(name)
    else:  # SKIP
        color = Colors.YELLOW
        symbol = "○"
        test_results["skipped"].append(name)

    print(f"{color}{symbol} {name}{Colors.RESET}")
    if message:
        print(f"  {message}")


def test_environment():
    """Test 1: Environment Setup"""
    print_header("1. ENVIRONMENT SETUP")

    # Check .env file
    if os.path.exists(".env"):
        print_test(".env file exists", "PASS")
    else:
        print_test(".env file exists", "FAIL", "Create .env from .env.example")
        return False

    # Check NOTION_TOKEN
    from dotenv import load_dotenv
    load_dotenv()
    token = os.getenv('NOTION_TOKEN')

    if token:
        if token.startswith('ntn_') or token.startswith('secret_'):
            print_test("NOTION_TOKEN format", "PASS")
        else:
            print_test("NOTION_TOKEN format", "FAIL", "Token should start with 'ntn_' or 'secret_'")
    else:
        print_test("NOTION_TOKEN exists", "FAIL", "Add NOTION_TOKEN to .env")
        return False

    # Check template file
    template_path = "IC Template.docx"
    if os.path.exists(template_path):
        print_test("Template file exists", "PASS")
    else:
        print_test("Template file exists", "SKIP", "Template not found, will use default")

    # Check output directory
    if os.path.exists("outputs"):
        print_test("Output directory exists", "PASS")
    else:
        os.makedirs("outputs", exist_ok=True)
        print_test("Output directory created", "PASS")

    return True


def test_url_parsing():
    """Test 2: URL Parsing"""
    print_header("2. URL PARSING")

    test_urls = [
        # (input, expected_id)
        ("https://www.notion.so/workspace/Page-Title-274e531aede680e9a2a7c31931ae355e",
         "274e531aede680e9a2a7c31931ae355e"),
        ("https://notion.so/274e531aede680e9a2a7c31931ae355e",
         "274e531aede680e9a2a7c31931ae355e"),
        ("274e531aede680e9a2a7c31931ae355e",
         "274e531aede680e9a2a7c31931ae355e"),
        ("https://www.notion.so/workspace/Test-Page-abc123def456?source=copy_link",
         "abc123def456"),
    ]

    all_passed = True
    for url, expected in test_urls:
        page_id, _ = NotionClient.extract_page_id(url)
        if page_id == expected:
            print_test(f"Parse: {url[:50]}...", "PASS")
        else:
            print_test(f"Parse: {url[:50]}...", "FAIL",
                      f"Expected {expected}, got {page_id}")
            all_passed = False

    return all_passed


def test_configuration():
    """Test 3: Configuration Loading"""
    print_header("3. CONFIGURATION")

    # Test config file loading
    config = TemplateConfig("template_config.json")

    if config.config:
        print_test("Load template_config.json", "PASS")
    else:
        print_test("Load template_config.json", "FAIL")
        return False

    # Test template retrieval
    ic_template = config.get_template("ic_template")
    if ic_template and ic_template.get("font") == "Aptos":
        print_test("IC template configuration", "PASS")
    else:
        print_test("IC template configuration", "FAIL", "Check template_config.json")

    # Test style mappings
    styles_to_test = ["paragraph", "heading_1", "bulleted_list_item", "unknown_type"]
    for block_type in styles_to_test:
        style = config.get_style_for_block(block_type, "ic_template")
        if style:
            print_test(f"Style mapping: {block_type}", "PASS", f"→ {style}")
        else:
            print_test(f"Style mapping: {block_type}", "FAIL")

    return True


def test_notion_connection():
    """Test 4: Notion API Connection"""
    print_header("4. NOTION API CONNECTION")

    from dotenv import load_dotenv
    load_dotenv()
    token = os.getenv('NOTION_TOKEN')

    if not token:
        print_test("Notion connection", "SKIP", "No token available")
        return False

    try:
        client = NotionClient(token)

        # Try to search for any content
        from notion_client import Client
        notion = Client(auth=token)

        response = notion.search(filter={"property": "object", "value": "page"})

        if response:
            print_test("API authentication", "PASS")
            print_test("Search capability", "PASS",
                      f"Found {len(response.get('results', []))} pages")
        else:
            print_test("API authentication", "FAIL")
            return False

    except Exception as e:
        print_test("Notion connection", "FAIL", str(e))
        return False

    return True


def test_document_creation():
    """Test 5: Document Creation"""
    print_header("5. DOCUMENT CREATION")

    # Test blank document
    try:
        formatter = WordFormatter({"font": "Aptos", "formatting": {}})
        doc = formatter.create_document()

        if doc:
            print_test("Create blank document", "PASS")
        else:
            print_test("Create blank document", "FAIL")
            return False
    except Exception as e:
        print_test("Create blank document", "FAIL", str(e))
        return False

    # Test with template
    template_path = "IC Template.docx"
    if os.path.exists(template_path):
        try:
            doc = formatter.create_document(template_path)
            print_test("Create from template", "PASS")
        except Exception as e:
            print_test("Create from template", "FAIL", str(e))
    else:
        print_test("Create from template", "SKIP", "No template found")

    # Test saving document
    try:
        test_output = "outputs/test_document.docx"
        formatter.add_title(doc, "Test Document")
        doc.save(test_output)

        if os.path.exists(test_output):
            print_test("Save document", "PASS", test_output)

            # Verify it can be opened
            test_doc = Document(test_output)
            if len(test_doc.paragraphs) > 0:
                print_test("Verify saved document", "PASS")

            # Cleanup
            os.remove(test_output)
        else:
            print_test("Save document", "FAIL")
    except Exception as e:
        print_test("Save document", "FAIL", str(e))

    return True


def test_sample_conversion():
    """Test 6: Sample Page Conversion"""
    print_header("6. SAMPLE CONVERSION")

    print("\n📝 Create a test Notion page with:")
    print("  - Some text paragraphs")
    print("  - A heading or two")
    print("  - A bulleted list")
    print("  - Share it with your integration")

    test_url = input("\nEnter test page URL (or press Enter to skip): ").strip()

    if not test_url:
        print_test("Sample conversion", "SKIP", "No test URL provided")
        return True

    from dotenv import load_dotenv
    load_dotenv()
    token = os.getenv('NOTION_TOKEN')

    if not token:
        print_test("Sample conversion", "SKIP", "No token available")
        return False

    try:
        converter = NotionToWordConverter(
            token,
            template_path="IC Template.docx" if os.path.exists("IC Template.docx") else None,
            output_dir="outputs"
        )

        output_path = converter.convert_page(test_url)

        if os.path.exists(output_path):
            # Check file size
            file_size = os.path.getsize(output_path)
            print_test("Convert sample page", "PASS",
                      f"Created {output_path} ({file_size:,} bytes)")

            # Verify content
            doc = Document(output_path)
            if len(doc.paragraphs) > 0:
                print_test("Document has content", "PASS",
                          f"{len(doc.paragraphs)} paragraphs")
            else:
                print_test("Document has content", "FAIL", "No paragraphs found")

        else:
            print_test("Convert sample page", "FAIL", "Output file not created")

    except Exception as e:
        print_test("Convert sample page", "FAIL", str(e))
        return False

    return True


def test_error_handling():
    """Test 7: Error Handling"""
    print_header("7. ERROR HANDLING")

    from dotenv import load_dotenv
    load_dotenv()
    token = os.getenv('NOTION_TOKEN')

    if not token:
        print_test("Error handling", "SKIP", "No token for testing")
        return True

    # Test invalid page ID
    try:
        client = NotionClient(token)
        result = client.fetch_page_content("invalid_page_id_12345")
        print_test("Invalid page ID handling", "FAIL", "Should have raised error")
    except:
        print_test("Invalid page ID handling", "PASS", "Error caught correctly")

    # Test invalid template path
    try:
        formatter = WordFormatter({"font": "Arial"})
        doc = formatter.create_document("nonexistent_template.docx")
        print_test("Invalid template handling", "PASS", "Fallback to blank document")
    except:
        print_test("Invalid template handling", "FAIL", "Should handle gracefully")

    return True


def test_cli_interface():
    """Test 8: CLI Interface"""
    print_header("8. COMMAND-LINE INTERFACE")

    # Test help
    import subprocess

    try:
        result = subprocess.run(
            ["python", "notion2word.py", "--help"],
            capture_output=True,
            text=True
        )

        if "Convert Notion pages" in result.stdout:
            print_test("CLI help command", "PASS")
        else:
            print_test("CLI help command", "FAIL", "Help text not found")
    except Exception as e:
        print_test("CLI help command", "FAIL", str(e))

    # Test without arguments (should prompt)
    # Skip interactive test
    print_test("CLI interactive mode", "SKIP", "Requires manual testing")

    return True


def run_performance_test():
    """Test 9: Performance Benchmarks"""
    print_header("9. PERFORMANCE")

    print("\nBenchmark requires a test page URL.")
    test_url = input("Enter page URL for benchmarking (or Enter to skip): ").strip()

    if not test_url:
        print_test("Performance benchmark", "SKIP", "No URL provided")
        return True

    from dotenv import load_dotenv
    load_dotenv()
    token = os.getenv('NOTION_TOKEN')

    if not token:
        print_test("Performance benchmark", "SKIP", "No token")
        return False

    try:
        converter = NotionToWordConverter(token)

        # Measure conversion time
        start_time = time.time()
        output_path = converter.convert_page(test_url)
        end_time = time.time()

        duration = end_time - start_time

        if duration < 5:
            print_test("Conversion speed", "PASS", f"Completed in {duration:.2f} seconds")
        elif duration < 10:
            print_test("Conversion speed", "PASS", f"Completed in {duration:.2f} seconds (acceptable)")
        else:
            print_test("Conversion speed", "FAIL", f"Took {duration:.2f} seconds (too slow)")

    except Exception as e:
        print_test("Performance benchmark", "FAIL", str(e))
        return False

    return True


def print_summary():
    """Print test summary"""
    print_header("TEST SUMMARY")

    total = len(test_results["passed"]) + len(test_results["failed"]) + len(test_results["skipped"])

    print(f"\n{Colors.GREEN}✓ Passed: {len(test_results['passed'])}{Colors.RESET}")
    print(f"{Colors.RED}✗ Failed: {len(test_results['failed'])}{Colors.RESET}")
    print(f"{Colors.YELLOW}○ Skipped: {len(test_results['skipped'])}{Colors.RESET}")
    print(f"\nTotal: {total} tests")

    if test_results["failed"]:
        print(f"\n{Colors.RED}Failed tests:{Colors.RESET}")
        for test in test_results["failed"]:
            print(f"  - {test}")

    if len(test_results["failed"]) == 0:
        print(f"\n{Colors.GREEN}{Colors.BOLD}All tests passed! 🎉{Colors.RESET}")
        return True
    else:
        print(f"\n{Colors.RED}{Colors.BOLD}Some tests failed. Please review.{Colors.RESET}")
        return False


def main():
    """Run all tests"""
    print(f"{Colors.BOLD}")
    print("╔" + "═" * 58 + "╗")
    print("║" + " NOTION TO WORD CONVERTER - TEST SUITE".center(58) + "║")
    print("╚" + "═" * 58 + "╝")
    print(f"{Colors.RESET}")

    # Run test suites
    test_environment()
    test_url_parsing()
    test_configuration()
    test_notion_connection()
    test_document_creation()
    test_sample_conversion()
    test_error_handling()
    test_cli_interface()
    run_performance_test()

    # Print summary
    return print_summary()


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)